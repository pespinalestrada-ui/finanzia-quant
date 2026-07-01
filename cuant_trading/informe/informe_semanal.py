"""
informe_semanal — informe .docx de tu watchlist en un clic (o programado).

Contenido: señales del sistema (score del Veredicto técnico) para cada valor de tu
watchlist + riesgo de la cesta (VaR/CVaR/drawdown) + titulares recientes. Rápido
(sin Prophet ni FinBERT) → tarda segundos. Guarda Informe_Semanal_AAAA-MM-DD.docx
en la raíz del proyecto. No es recomendación de inversión.

Uso:
    python informe_semanal.py
    python informe_semanal.py AAPL MSFT NVDA KO
"""
import sys
from datetime import date
from pathlib import Path
try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass
import warnings
warnings.filterwarnings("ignore")

_SUITE = Path(__file__).resolve().parents[1]
_PROJ = _SUITE.parent
for p in (_SUITE / "signal_engine", _SUITE / "risk_metrics", _SUITE / "sentiment",
          _SUITE / "veredicto_backtest"):
    sys.path.insert(0, str(p))


def _watchlist():
    f = _PROJ / "watchlist.txt"
    if f.exists():
        txt = f.read_text(encoding="utf-8").strip()
        if txt:
            return [t.strip().upper() for t in txt.replace("\n", ",").split(",") if t.strip()]
    return ["AAPL", "MSFT", "NVDA", "GOOGL", "AMZN", "META", "JPM", "XOM", "KO", "SAB.MC"]


def generar(tickers=None):
    """Genera el informe. Devuelve la ruta del .docx."""
    import docx
    from docx.shared import Pt, RGBColor
    import signal_engine as SE
    import risk_metrics as RM
    tickers = tickers or _watchlist()

    señales = SE.generar(tickers, 0.35)
    try:
        riesgo_tabla, corr, _curva, meta_r = RM.analizar(tickers, "1y", 0.95)
    except Exception:
        riesgo_tabla, meta_r = None, {}
    noticias = {}
    try:
        from news_feeds import obtener_noticias
        for tk in tickers[:5]:
            ns = obtener_noticias(tk, 3)
            if ns:
                noticias[tk] = [n["titular"][:110] for n in ns[:3]]
    except Exception:
        pass

    d = docx.Document()
    st = d.styles["Normal"]; st.font.name = "Arial"; st.font.size = Pt(10)
    h = d.add_heading("FinanzIA — Informe semanal", level=0)
    d.add_paragraph(f"Fecha: {date.today().isoformat()} · Watchlist: {', '.join(tickers)}")
    d.add_paragraph("Herramienta de análisis y educación. No es recomendación de inversión."
                    ).italic = True

    d.add_heading("1. Señales de la semana (score técnico del Veredicto)", level=1)
    if señales is not None and len(señales):
        t = d.add_table(rows=1, cols=4); t.style = "Light Grid Accent 1"
        for i, c in enumerate(["Ticker", "Precio", "Score", "Señal"]):
            t.rows[0].cells[i].text = c
        for _, r in señales.iterrows():
            row = t.add_row().cells
            row[0].text = str(r["Ticker"]); row[1].text = str(r["Precio"])
            row[2].text = str(r["Score"]); row[3].text = str(r["Señal"])
        comprar = señales[señales["Señal"].str.contains("COMPRAR")]["Ticker"].tolist()
        vender = señales[señales["Señal"].str.contains("VENDER")]["Ticker"].tolist()
        d.add_paragraph(f"En señal de compra: {', '.join(comprar) if comprar else 'ninguno'} · "
                        f"en señal de venta: {', '.join(vender) if vender else 'ninguno'}.")
    else:
        d.add_paragraph("Sin datos de señales esta semana.")

    d.add_heading("2. Riesgo de la cesta", level=1)
    if riesgo_tabla is not None and len(riesgo_tabla):
        t = d.add_table(rows=1, cols=len(riesgo_tabla.columns)); t.style = "Light Grid Accent 1"
        for i, c in enumerate(riesgo_tabla.columns):
            t.rows[0].cells[i].text = str(c)
        for _, r in riesgo_tabla.iterrows():
            row = t.add_row().cells
            for i, c in enumerate(riesgo_tabla.columns):
                row[i].text = str(r[c])
        cm = meta_r.get("corr_media")
        if cm is not None:
            d.add_paragraph(f"Correlación media entre valores: {cm:.2f} "
                            f"({'buena diversificación' if cm < 0.4 else 'poca diversificación'}).")
    else:
        d.add_paragraph("Riesgo no disponible.")

    d.add_heading("3. Titulares recientes", level=1)
    if noticias:
        for tk, ts in noticias.items():
            d.add_paragraph(tk + ":", style="Intense Quote")
            for x in ts:
                d.add_paragraph("• " + x)
    else:
        d.add_paragraph("Sin noticias disponibles ahora mismo.")

    d.add_paragraph()
    d.add_paragraph("Recuerda: el valor está en controlar el riesgo, no en adivinar el precio. "
                    "Practica en paper antes de arriesgar dinero real.").italic = True
    out = _PROJ / f"Informe_Semanal_{date.today().isoformat()}.docx"
    d.save(str(out))
    return str(out)


if __name__ == "__main__":
    tks = [t.upper() for t in sys.argv[1:]] or None
    ruta = generar(tks)
    print(f"Informe generado: {ruta}")
